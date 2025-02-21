import os
import fitz
from autoblock import Autoblock
from pprint import pprint
from docx import Document

class DOC_reader:
    def __init__(self, name, win):
        self.name = name
        self.win = win
        self.location = os.path.join(self.win.folder, name)
        self.doc = Document(self.location)
        self.page=None
        self.tables=None
        self.solve_auto()
        self.doc.save(self.location)



    def Doc_text(self):
        self.doc.add_heading("Test",0)
        for p in self.doc.paragraphs:
           print(p.text)

    def print(self):
        text=""
        for page_num in range(len(self.doc)):
           self.page = self.doc[page_num]
           text += self.page.get_text()
           lines=text.split("\n")
           fil_lines=[]
           for line in lines:
               if line!=" ":
                   fil_lines.append(line)
           ind=0
           print(fil_lines)
           while ind<len(fil_lines):
               if fil_lines[ind]=="+ ":
                   block=Autoblock(fil_lines[ind] ,fil_lines[ind+1:(ind+21)])
               ind+=1

    def find_table(self):
        page=self.doc[3]
        tabs=page.find_tables()
        print(f"{len(tabs.tables)} found on {page}")

        if tabs.tables:
            pprint(tabs[0].extract())

    def solve_auto(self):
        for table in self.doc.tables:
            if table.cell(0,0).text=="+":
                for i in range(1, len(table.rows)):
                    for g in range(1, len(table.columns)):
                        table.cell(i,g).text=str(int(table.cell(0,g).text)+int(table.cell(i,0).text))
            elif table.cell(0,0).text=="-":
                for i in range(1, len(table.rows)):
                    for g in range(1, len(table.columns)):
                        table.cell(i,g).text=str(int(table.cell(0,g).text)-int(table.cell(i,0).text))
            elif table.cell(0,0).text=="x":
                if table.cell(1,0).text != " ":
                    for i in range(1, len(table.rows)):
                        for g in range(1, len(table.columns)):
                            table.cell(i,g).text=str(int(table.cell(0,g).text)*int(table.cell(i,0).text))
