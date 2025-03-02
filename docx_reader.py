import os
import tkinter
from docx import Document
from tkinter.ttk import *
import random

class DOC_reader:
    def __init__(self, name, win, create=False, answer=False):
        self.name = name
        self.win = win
        self.location = os.path.join(self.win.folder, name)
        self.doc = Document(self.location)
        self.page=None
        self.create=create
        self.progres=None
        self.prog_var=tkinter.DoubleVar()
        self.answer=answer
        self.max=[]
        self.min=[]
        self.types=[]
        self.types_set=[]
        self.inverse=[]
        self.inverse_buttons=[]


        if self.create is False:
            self.progress()
            self.strip_errors()

        self.tables=self.get_tables()

        if self.create is True:
            self.input_min_max(self.win.width/4)
            button_copy=tkinter.Button(text="Create", background="green", command=self.create_setup)
            self.win.canvas.create_window(600,250,window=button_copy)

        if self.answer is False:
            button_answer=tkinter.Button(text="Answer", background="green", command=self.solve_auto)
            self.win.canvas.create_window((self.win.width/3)*2,100,window=button_answer)
            button_create=tkinter.Button(text="Create", background="green", command=self.create_auto)
            self.win.canvas.create_window(300,100,window=button_create)

        if self.answer is True:
            self.solve_auto()

    def get_tables(self):
        tables=[]
        for table in self.doc.tables:
            if table.cell(0,0).text=="+" or table.cell(0,0).text=="-" or table.cell(0,0).text=="x":
                tables.append(table)
        return tables

    def solve_addition(self,table,k):
        for i in range(1, len(table.rows)):
            for g in range(1, len(table.columns)):
                if table.cell(i,0).text != "" and table.cell(0,g).text != "":
                   table.cell(i,g).text=str(int(table.cell(0,g).text)+int(table.cell(i,0).text))
                   k=self.progress_update(k)
        return k

    def solve_subtraction(self, table,k):
        for i in range(1, len(table.rows)):
            for g in range(1, len(table.columns)):
                if table.cell(i,0).text != "" and table.cell(0,g).text != "":
                    table.cell(i,g).text=str(float(table.cell(0,g).text)-float(table.cell(i,0).text))
                    k=self.progress_update(k)
        return k

    def solve_multiplication(self, table,k):
        for i in range(1, len(table.rows)):
            for g in range(1, len(table.columns)):
                if table.cell(i,0).text != "" and table.cell(0,g).text != "":
                    table.cell(i,g).text=str(float(table.cell(0,g).text)*float(table.cell(i,0).text))
                    k=self.progress_update(k)
        return k

    def solve_division(self, table,k):
        for i in range(1, len(table.rows)):
            for g in range(1, len(table.columns)):
                if table.cell(i,g).text != " " and table.cell(i,g).text != "":
                    table.cell(i,0).text=str(float(table.cell(i,g).text)/float(table.cell(0,g).text))
                    k=self.progress_update(k,0.5)
        for i in range(1, len(table.rows)):
            for g in range(1, len(table.columns)):
                if table.cell(i,0).text != " " and table.cell(i,0).text != "":
                    table.cell(i,g).text=str(float(table.cell(0,g).text)*float(table.cell(i,0).text))
                    k=self.progress_update(k,0.5)
        return k

    def solve_auto(self):
        k=0.0
        self.progress()
        for table in self.doc.tables:
            if table.cell(0,0).text=="+":
                k=self.solve_addition(table,k)
            elif table.cell(0,0).text=="-":
                k=self.solve_subtraction(table,k)
            elif table.cell(0,0).text=="x":
                if table.cell(1,0).text != "":
                    k=self.solve_multiplication(table,k)
                else:
                    k=self.solve_division(table,k)
        self.doc.save(os.path.join(self.win.folder, self.name[:-5]+"-answers"+self.name[-5:]))
        self.win.popup_text("Answers Complete")

    def create_auto(self, min, max,tables, symbol):
        k=0
        self.progress()
        for table in tables:
            options=["+","-","x"]
            if symbol=="mixed":
                z=options[random.randint(0,2)]
            else:
                z=symbol
            table.cell(0,0).text=z
            for i in range(1, len(table.columns)):
                table.cell(0,i).text=str(random.randint(min, max))
                k=self.progress_update(k)+15
            for g in range(1,  len(table.rows)):
                table.cell(g,0).text=str(random.randint(min, max))
                k=self.progress_update(k)+15


    def strip_errors(self):
        k=0
        self.progress()
        for table in self.tables:
            zero_zero=table.cell(0,0).text
            for i in range(len(table.rows)):
                for g in range(len(table.columns)):
                    table.cell(i,g).text="".join(c for c in table.cell(i,g).text if c.isdigit() or c=="-")
                    k=self.progress_update(k)
            table.cell(0,0).text=zero_zero
        self.doc.save(self.location)


    def progress(self):
        self.progres = Progressbar(self.win.canvas, orient = 'horizontal',variable=self.prog_var, maximum = (len(self.doc.tables)*110), mode = 'determinate')
        self.progres.place(x=self.win.width/4,y=480,width=self.win.width/2)

    def progress_update(self, k,x=1.0):
        self.prog_var.set(k)
        self.win.root.update_idletasks()
        k+=x
        return k

    def input_min_max(self, x):
        i=0
        options=["+","-","x","mixed"]
        while i<3:
            self.types_set.append(tkinter.StringVar(value=options[0]))

            max=tkinter.Label(self.win.root, text=f"Section {i+1} Max")
            self.win.canvas.create_window(x*(i+1),330, window=max)

            self.max.append(tkinter.Entry(self.win.root))
            self.win.canvas.create_window(x*(i+1), 360, window=self.max[i])

            min=tkinter.Label(self.win.root, text=f"Section {i+1} Min")
            self.win.canvas.create_window(x*(i+1),390, window=min)

            self.min.append(tkinter.Entry(self.win.root))
            self.win.canvas.create_window(x*(i+1), 420, window=self.min[i])

            self.types.append(tkinter.OptionMenu(self.win.root,self.types_set[i], *options))
            self.win.canvas.create_window(x*(i+1),450,window=self.types[i])

            self.inverse.append(tkinter.IntVar())
            self.inverse_buttons.append(tkinter.Checkbutton(self.win.root, text = "Inverse", variable = self.inverse[i], onvalue = 1, offvalue = 0))
            self.win.canvas.create_window(x*(i+1),300,window=self.inverse_buttons[i])

            i+=1

    def create_setup(self):
        k=0
        while k<3:
            x=self.types_set[k].get()
            if self.inverse[k].get()==0:
                self.create_auto(int(self.min[k].get()),int(self.max[k].get()),self.tables[int(len(self.tables)/3*k):int(len(self.tables)/3*(k+1))], x)
                k+=1
            elif self.inverse[k].get()==1:
                self.create_inverse_auto(int(self.min[k].get()),int(self.max[k].get()),self.tables[int(len(self.tables)/3*k):int(len(self.tables)/3*(k+1))], x)
                k+=1
        self.doc.save(os.path.join(self.win.folder, self.name[:-5]+"-random"+self.name[-5:]))
        self.win.popup_text("Student Copy Complete")
        self.__init__(self.name[:-5]+"-random"+self.name[-5:], self.win, answer=True)

    def create_inverse_auto(self, min, max,tables, symbol):
            k=0
            self.progress()
            for table in tables:
                options=["+","-","x"]
                if symbol=="mixed":
                    z=options[random.randint(0,2)]
                else:
                    z=symbol
                g=1
                table.cell(0,0).text=z
                for i in range(1, len(table.columns)):
                    table.cell(0,i).text=str(random.randint(min, max))
                    k=self.progress_update(k)+15
                while g <len(table.rows):
                    y=random.randint(1,10)
                    if z=="x":
                        table.cell(g,0).text=""
                        table.cell(g,y).text=str(random.randrange(0, y*12,int(table.cell(0,y).text)))
                        g=g+1
                        k=self.progress_update(k)+15
                    else:
                        table.cell(g,0).text=""
                        table.cell(g,y).text=str(random.randint(min, max))
                        g=g+1
                        k=self.progress_update(k)+15
